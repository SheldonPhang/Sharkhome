import os
import docx
import datetime
def generate_report(scan_results, oa_type):
    report_path = "report"
    if not os.path.exists(report_path):
        os.makedirs(report_path)
    
    now = datetime.datetime.now().strftime("%Y%m%d%H%M")
    report_file = os.path.join(report_path, f"{oa_type}_scan_report_{now}.docx")
    doc = docx.Document()

    doc.add_heading(f"{oa_type} 漏洞扫描报告", 0)

    for item in scan_results:
        doc.add_paragraph(str(item), style="BodyText")

    doc.save(report_file)
    return report_file